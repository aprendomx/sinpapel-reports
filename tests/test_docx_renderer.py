from __future__ import annotations

import io

from docx import Document as DocxDocument

from sinpapel_reports.services.docx_renderer import DocxRenderer


def _docx_with_placeholder(path: str) -> None:
    doc = DocxDocument()
    doc.add_paragraph("Hola {{ nombre }}")
    doc.save(path)


def test_render_fills_placeholder(tmp_path):
    tpl = tmp_path / "t.docx"
    _docx_with_placeholder(str(tpl))
    out = DocxRenderer.render(str(tpl), {"nombre": "MUNDO"})
    rendered = DocxDocument(io.BytesIO(out))
    text = "\n".join(p.text for p in rendered.paragraphs)
    assert "Hola MUNDO" in text
